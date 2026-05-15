"""
PDF to Word Tool - Converts PDFs to DOCX using pdf2docx Layout Mode.

This tool follows the same visual pattern as preview-oriented PDF tools:
left-side controls and right-side page preview. Conversion runs in a
background thread and preserves the shared output actions.
"""

import os
import queue
import tempfile
import threading
from io import BytesIO
from pathlib import Path
from typing import Any, Dict, List, Optional

import fitz
import tkinter as tk
from docx import Document
from docx.shared import Inches
from pdf2docx import Converter
from PIL import Image, ImageTk
from tkinter import filedialog, messagebox, ttk

from ...base.tool import BaseTool
from ...ui.components import FileListWidget, OutputActions
from ...utils.file_ops import get_default_save_dir
from ...utils.settings import get_setting, set_setting


class PDFToWordTool(BaseTool):
    """GUI tool for converting PDF files to Word DOCX files."""

    name: str = "PDF to Word"
    icon: str = "[W]"
    MODES = ("Preserve Layout", "Text Only", "Page Images")

    def __init__(self) -> None:
        self.queue: queue.Queue = queue.Queue()
        self.doc: Optional[fitz.Document] = None
        self.current_pdf_path: Optional[str] = None
        self.total_pages = 0
        self.preview_img: Optional[ImageTk.PhotoImage] = None

    def render(self, parent: ttk.Frame) -> None:
        paned = ttk.PanedWindow(parent, orient=tk.HORIZONTAL)
        paned.pack(fill="both", expand=True, pady=5)

        left_frame = ttk.Frame(paned)
        paned.add(left_frame, weight=1)

        self.file_list = FileListWidget(
            left_frame,
            label="Source PDFs (click to preview)",
            filetypes=[("PDF Files", "*.pdf")],
            show_ordering=False,
            display_formatter=lambda path: os.path.basename(path),
        )
        self.file_list.pack(fill="both", expand=True, pady=5)
        self.file_list.bind_select(self._on_file_select)

        settings_frame = ttk.LabelFrame(left_frame, text="Conversion Settings", padding=10)
        settings_frame.pack(fill="x", pady=5)

        ttk.Label(settings_frame, text="Mode:").grid(
            row=0, column=0, sticky="w", padx=(0, 10)
        )
        self.mode_var = tk.StringVar(
            value=get_setting("pdf2word.mode", "Preserve Layout")
        )
        ttk.Combobox(
            settings_frame,
            textvariable=self.mode_var,
            values=self.MODES,
            state="readonly",
            width=18,
        ).grid(
            row=0, column=1, sticky="w", padx=(0, 25)
        )

        ttk.Label(settings_frame, text="Page Range:").grid(
            row=1, column=0, sticky="w", padx=(0, 10), pady=(8, 0)
        )
        self.range_entry = ttk.Entry(settings_frame)
        self.range_entry.grid(row=1, column=1, columnspan=2, sticky="ew", pady=(8, 0))
        settings_frame.columnconfigure(2, weight=1)

        ttk.Label(
            settings_frame,
            text="Blank = all pages. Page Images avoids blank pages but is not editable.",
            foreground="gray",
        ).grid(row=2, column=0, columnspan=3, sticky="w", pady=(8, 0))

        self.preflight_lbl = ttk.Label(
            settings_frame, text="Select a PDF to preview and preflight.", foreground="gray"
        )
        self.preflight_lbl.grid(row=3, column=0, columnspan=3, sticky="w", pady=(8, 0))

        out_frame = ttk.LabelFrame(left_frame, text="Output Folder", padding=10)
        out_frame.pack(fill="x", pady=5)
        self.output_entry = ttk.Entry(out_frame)
        self.output_entry.pack(side="left", fill="x", expand=True, padx=(0, 5))
        self.output_entry.insert(
            0, get_setting("pdf2word.output_dir", get_default_save_dir("Word"))
        )
        ttk.Button(out_frame, text="Browse", command=self._browse_output).pack(
            side="right"
        )

        self.btn = ttk.Button(left_frame, text="Convert to Word", command=self.execute)
        self.btn.pack(pady=15, fill="x")

        self.progress = ttk.Progressbar(left_frame, mode="determinate")
        self.progress.pack(fill="x")
        self.status_lbl = ttk.Label(left_frame, text="Ready.")
        self.status_lbl.pack(anchor="w")

        self.output_actions = OutputActions(left_frame)
        self.output_actions.pack(anchor="w", pady=(8, 0))

        right_frame = ttk.LabelFrame(paned, text="PDF Preview", padding=10)
        paned.add(right_frame, weight=3)

        nav_frame = ttk.Frame(right_frame)
        nav_frame.pack(fill="x", pady=(0, 5))

        self.page_lbl = ttk.Label(nav_frame, text="Page: -/-")
        self.page_lbl.pack(side="left")

        self.preview_var = tk.IntVar(value=1)
        self.preview_scale = ttk.Scale(
            nav_frame,
            from_=1,
            to=1,
            variable=self.preview_var,
            command=self._on_preview_change,
        )
        self.preview_scale.pack(side="right", fill="x", expand=True, padx=10)

        self.preview_lbl = ttk.Label(right_frame, text="No PDF loaded", anchor="center")
        self.preview_lbl.pack(fill="both", expand=True)
        self.preview_lbl.bind("<MouseWheel>", self._on_mouse_wheel)
        self.preview_lbl.bind("<Button-4>", self._on_mouse_wheel)
        self.preview_lbl.bind("<Button-5>", self._on_mouse_wheel)

        self._process_queue()

    def _browse_output(self) -> None:
        path = filedialog.askdirectory(title="Select Output Folder")
        if path:
            self.output_entry.delete(0, tk.END)
            self.output_entry.insert(0, path)

    def _on_file_select(self, event: Any) -> None:
        selected = self.file_list.get_selected_file()
        if selected and selected != self.current_pdf_path:
            self._load_pdf(selected)

    def _load_pdf(self, path: str) -> None:
        try:
            if self.doc:
                self.doc.close()
            doc = fitz.open(path)
            if doc.needs_pass:
                doc.close()
                raise ValueError("This PDF is encrypted and cannot be previewed.")
            if doc.page_count == 0:
                doc.close()
                raise ValueError("This PDF has no pages.")

            self.doc = doc
            self.current_pdf_path = path
            self.total_pages = doc.page_count
            self.preview_scale.config(to=max(1, self.total_pages))
            self.preview_var.set(1)
            self._update_preview(0)
            self._update_preflight_label(path)
        except Exception as exc:
            self.doc = None
            self.current_pdf_path = None
            self.total_pages = 0
            self.page_lbl.config(text="Page: -/-")
            self.preview_lbl.config(image="", text="No PDF loaded")
            self.preflight_lbl.config(text=f"Preview unavailable: {exc}")
            messagebox.showerror("Error", f"Failed to load PDF preview: {exc}")

    def _update_preflight_label(self, path: str) -> None:
        try:
            report = self.preflight_pdf(path, self.range_entry.get().strip())
            message = (
                f"{report['page_count']} pages, "
                f"{report['selected_page_count']} selected. "
            )
            if report["image_only"]:
                message += "Looks image-only/scanned; OCR is planned for a later phase."
            else:
                message += "Text detected."
            self.preflight_lbl.config(text=message)
        except Exception as exc:
            self.preflight_lbl.config(text=f"Preflight warning: {exc}")

    def _on_preview_change(self, value: Any) -> None:
        page_num = int(float(value)) - 1
        self._update_preview(page_num)

    def _on_mouse_wheel(self, event: Any) -> None:
        if self.total_pages == 0:
            return
        delta = 1 if (getattr(event, "num", None) == 5 or event.delta < 0) else -1
        new_page = max(1, min(self.total_pages, self.preview_var.get() + delta))
        self.preview_var.set(new_page)
        self._update_preview(new_page - 1)

    def _update_preview(self, page_num: int) -> None:
        if not self.doc or page_num < 0 or page_num >= self.total_pages:
            return
        self.page_lbl.config(text=f"Page: {page_num + 1}/{self.total_pages}")
        try:
            page = self.doc[page_num]
            pix = page.get_pixmap(dpi=100)
            mode = "RGBA" if pix.alpha else "RGB"
            img = Image.frombytes(mode, [pix.width, pix.height], pix.samples)

            max_height = 560
            if img.height > max_height:
                ratio = max_height / img.height
                img = img.resize(
                    (int(img.width * ratio), max_height), Image.Resampling.LANCZOS
                )

            tk_img = ImageTk.PhotoImage(img)
            self.preview_lbl.config(image=tk_img, text="")
            self.preview_img = tk_img
        except Exception as exc:
            self.preview_lbl.config(image="", text=f"Preview failed: {exc}")

    def _process_queue(self) -> None:
        try:
            while True:
                msg_type, data = self.queue.get_nowait()
                if msg_type == "progress":
                    pct, message = data
                    self.progress["value"] = pct
                    self.status_lbl.config(text=message)
                elif msg_type == "success":
                    self.btn.config(state="normal")
                    self.progress["value"] = 100
                    output_dir = data["output_dir"]
                    self.output_actions.set_path(output_dir)
                    self.status_lbl.config(text=data["message"])
                    messagebox.showinfo("Success", data["message"])
                elif msg_type == "error":
                    self.btn.config(state="normal")
                    self.progress["value"] = 0
                    self.status_lbl.config(text="Error occurred.")
                    messagebox.showerror("Error", data)
        except queue.Empty:
            pass

        if hasattr(self, "status_lbl") and self.status_lbl.winfo_exists():
            self.status_lbl.after(100, self._process_queue)

    def execute(self, params: Optional[Dict[str, Any]] = None) -> None:
        files = self.file_list.get_files()
        if not files:
            messagebox.showwarning("Warning", "No files selected!")
            return

        output_dir = self.output_entry.get()
        if not output_dir:
            messagebox.showwarning("Warning", "Please choose an output folder.")
            return

        range_text = self.range_entry.get().strip()
        mode = self.mode_var.get()
        preflight = self.preflight_files(files, range_text, mode)
        if preflight["fatal_errors"]:
            messagebox.showerror("Preflight Failed", "\n".join(preflight["fatal_errors"][:8]))
            return
        if preflight["warnings"]:
            messagebox.showwarning(
                "Preflight Warnings",
                "\n".join(preflight["warnings"][:8]),
            )

        set_setting("pdf2word.output_dir", output_dir)
        set_setting("pdf2word.mode", mode)

        self.btn.config(state="disabled")
        self.progress["value"] = 0
        self.output_actions.clear()
        self.status_lbl.config(text="Converting PDFs to Word...")

        threading.Thread(
            target=self._run_conversion,
            args=(files, output_dir, range_text, mode),
        ).start()

    def _run_conversion(
        self, files: List[str], output_dir: str, range_text: str, mode: str
    ) -> None:
        results = {"success": 0, "failed": 0, "skipped": 0, "errors": []}
        total = len(files)

        try:
            os.makedirs(output_dir, exist_ok=True)

            for idx, input_path in enumerate(files, start=1):
                pct = ((idx - 1) / total) * 100
                self.queue.put(
                    (
                        "progress",
                        (pct, f"Converting {idx}/{total}: {os.path.basename(input_path)}"),
                    )
                )

                try:
                    if not os.path.exists(input_path):
                        results["skipped"] += 1
                        results["errors"].append(f"File not found: {input_path}")
                        continue

                    output_path = self._output_path(input_path, output_dir)
                    self.convert_pdf_to_docx(input_path, output_path, range_text, mode)
                    results["success"] += 1
                except Exception as exc:
                    results["failed"] += 1
                    results["errors"].append(
                        f"{os.path.basename(input_path)}: {str(exc)}"
                    )

            self.queue.put(("progress", (100, "Conversion complete.")))
            summary = (
                f"Word conversion complete ({mode}). Success: {results['success']}, "
                f"failed: {results['failed']}, skipped: {results['skipped']}."
            )
            if results["errors"]:
                summary += "\n" + "\n".join(results["errors"][:5])
            self.queue.put(("success", {"message": summary, "output_dir": output_dir}))
        except Exception as exc:
            self.queue.put(("error", str(exc)))

    @classmethod
    def convert_pdf_to_docx(
        cls,
        input_path: str,
        output_path: str,
        range_text: str = "",
        mode: str = "Preserve Layout",
    ) -> None:
        """
        Converts a PDF to DOCX using the selected quality mode.

        Preserve Layout produces the most editable layout, while Page Images is
        the visual-fidelity fallback for formulas, scanned pages, and complex
        PDFs that do not convert cleanly into editable Word content.
        """
        cls.preflight_pdf(input_path, range_text)
        page_indices = cls.parse_page_range(range_text, input_path)

        if mode == "Text Only":
            cls._convert_text_only(input_path, output_path, page_indices)
            return
        if mode == "Page Images":
            cls._convert_page_images(input_path, output_path, page_indices)
            return

        temp_path = ""
        source_path = input_path

        try:
            if page_indices is not None:
                temp_path = cls._create_selected_pages_pdf(input_path, page_indices)
                source_path = temp_path

            converter = Converter(source_path)
            try:
                converter.convert(output_path)
            finally:
                converter.close()
        finally:
            if temp_path and os.path.exists(temp_path):
                try:
                    os.remove(temp_path)
                except OSError:
                    pass

    @classmethod
    def preflight_files(
        cls, files: List[str], range_text: str, mode: str = "Preserve Layout"
    ) -> Dict[str, List[str]]:
        warnings: List[str] = []
        fatal_errors: List[str] = []

        for input_path in files:
            try:
                report = cls.preflight_pdf(input_path, range_text)
                if report["image_only"] and mode != "Page Images":
                    warnings.append(
                        f"{os.path.basename(input_path)} looks image-only/scanned. "
                        "Use Page Images if Preserve Layout creates blank pages. OCR is not included yet."
                    )
            except Exception as exc:
                fatal_errors.append(f"{os.path.basename(input_path)}: {exc}")

        return {"warnings": warnings, "fatal_errors": fatal_errors}

    @staticmethod
    def _selected_pages(doc: fitz.Document, page_indices: Optional[List[int]]) -> List[int]:
        return page_indices if page_indices is not None else list(range(doc.page_count))

    @classmethod
    def _convert_text_only(
        cls, input_path: str, output_path: str, page_indices: Optional[List[int]]
    ) -> None:
        document = Document()
        document.add_heading(Path(input_path).stem, level=1)

        with fitz.open(input_path) as doc:
            selected = cls._selected_pages(doc, page_indices)
            for count, page_index in enumerate(selected):
                if count:
                    document.add_page_break()
                document.add_heading(f"Page {page_index + 1}", level=2)
                text = doc[page_index].get_text("text").strip()
                if text:
                    for block in text.split("\n\n"):
                        clean = " ".join(line.strip() for line in block.splitlines() if line.strip())
                        if clean:
                            document.add_paragraph(clean)
                else:
                    document.add_paragraph("[No extractable text on this page.]")

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

    @classmethod
    def _convert_page_images(
        cls, input_path: str, output_path: str, page_indices: Optional[List[int]]
    ) -> None:
        document = Document()
        section = document.sections[0]
        usable_width = section.page_width - section.left_margin - section.right_margin
        image_width = Inches(usable_width / 914400)

        with fitz.open(input_path) as doc:
            selected = cls._selected_pages(doc, page_indices)
            for count, page_index in enumerate(selected):
                if count:
                    document.add_page_break()
                pix = doc[page_index].get_pixmap(dpi=150, alpha=False)
                image = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)
                buffer = BytesIO()
                image.save(buffer, format="PNG")
                buffer.seek(0)
                document.add_picture(buffer, width=image_width)

        Path(output_path).parent.mkdir(parents=True, exist_ok=True)
        document.save(output_path)

    @classmethod
    def preflight_pdf(cls, input_path: str, range_text: str = "") -> Dict[str, Any]:
        if not os.path.exists(input_path):
            raise FileNotFoundError(input_path)

        with fitz.open(input_path) as doc:
            if doc.needs_pass:
                raise ValueError("PDF is encrypted and cannot be converted.")
            if doc.page_count == 0:
                raise ValueError("PDF has no pages.")

        page_indices = cls.parse_page_range(range_text, input_path)
        with fitz.open(input_path) as doc:
            selected = page_indices if page_indices is not None else list(range(doc.page_count))
            text_pages = 0
            for page_index in selected[:10]:
                if doc[page_index].get_text("text").strip():
                    text_pages += 1

            return {
                "page_count": doc.page_count,
                "selected_page_count": len(selected),
                "text_pages_sampled": text_pages,
                "image_only": text_pages == 0,
            }

    @staticmethod
    def parse_page_range(range_text: str, input_path: str) -> Optional[List[int]]:
        if not range_text.strip():
            return None

        with fitz.open(input_path) as doc:
            total_pages = doc.page_count
            if doc.needs_pass:
                raise ValueError("PDF is encrypted and cannot be read.")

        page_indices: List[int] = []
        for raw_part in range_text.split(","):
            part = raw_part.strip()
            if not part:
                continue
            if "-" in part:
                start, end = map(int, part.split("-", 1))
            else:
                start = end = int(part)

            if start < 1 or end < 1 or start > end or end > total_pages:
                raise ValueError(
                    f"Page range {part} is out of bounds. PDF has {total_pages} pages."
                )
            page_indices.extend(range(start - 1, end))

        if not page_indices:
            raise ValueError("Page range is empty.")
        return page_indices

    @staticmethod
    def _create_selected_pages_pdf(input_path: str, page_indices: List[int]) -> str:
        with tempfile.NamedTemporaryFile(
            suffix=".pdf", prefix="pdf_to_word_", delete=False
        ) as temp_file:
            temp_path = temp_file.name

        with fitz.open(input_path) as source_doc:
            selected_doc = fitz.open()
            try:
                for page_index in page_indices:
                    selected_doc.insert_pdf(
                        source_doc, from_page=page_index, to_page=page_index
                    )
                selected_doc.save(temp_path)
            finally:
                selected_doc.close()

        return temp_path

    @staticmethod
    def _output_path(input_path: str, output_dir: str) -> str:
        output_path = Path(output_dir) / f"{Path(input_path).stem}.docx"
        if not output_path.exists():
            return str(output_path)

        counter = 2
        while True:
            candidate = output_path.with_name(f"{output_path.stem}_{counter}.docx")
            if not candidate.exists():
                return str(candidate)
            counter += 1
